import docx
from docx import Document
import unicodedata
import os
from time import sleep

circumstances_list=["AS TO THE FACTS","THE FACTS","I.	The Circumstances of the case","the circumstances of the case","I.  THE CIRCUMSTANCES OF THE CASE","I. THE CIRCUMSTANCES OF THE CASE","I. THE PARTICULAR CIRCUMSTANCES OF THE CASE","I.   Particular circumstances of the case","the CIRCUMSTANCES OF THE CASE"]
law_list=["Domestic law","I.  DOMESTIC LAW AND PRACTICE","I.   RELEVANT DOMESTIC LAW","II.   DOMESTIC LAW AND PRACTICE","II.     The relevant legislation","II. THE RELEVANT LEGISLATION","II.  RELEVANT DOMESTIC LAW AND BACKGROUND","II.  Relevant domestic law and practice","ii.	Relevant domestic law and practice","II.  RELEVANT","III.  relevant","II.  DOMESTIC LAW AND PRACTICE","II. RELEVANT","II.   Relevant","ii.	relevant domestic","II.  DOMESTIC LAW AND PRACTICE","II.	Relevant domestic law and practice","II.	Relevant domestic law","AS TO THE LAW","THE LAW","II.   RELEVANT DOMESTIC LAW AND PRACTICE"]

##########cut everything below fact sectiion
# directorycut="CaseLawcomplete"
# for filename in os.listdir(directorycut):
# # document=Document('CaseLawPDF141\\001-60876.docx')
#     try:
#        document = Document(directorycut + "\\" + filename)
#     except:
#         continue
#     file=filename.strip('docx')
#     newf=open("CaseLawcompletecut\\"+file+"txt",'w',encoding='utf-8')
#     for j in range(len(document.paragraphs)):
#         # print(document.paragraphs[j].text)
#         # print(j)
#         if all([x not in unicodedata.normalize("NFKD",document.paragraphs[j].text) for x in law_list]):
#                 # print(j)
#                 newf.write(document.paragraphs[j].text+"\n")
#         else:
#                 break

#######second cut: everything above fact section
# directorycut="CaseLawcompletecut"
# # for filename in os.listdir(directorycut):
# files=["001-57416.txt","001-57456.txt","001-57498.txt","001-57533.txt","001-57545.txt"]
# for filename in files:
#     file=open(directorycut+"\\"+filename,"r",encoding='utf-8')
#     lines = file.readlines()
#     newf=open("CaseLawcompletecutfacts\\"+filename,'w',encoding='utf-8')
#     j=0
#     # print(len(lines))
#     while j < len(lines)-1:
#         if any([x in unicodedata.normalize("NFKD", lines[j]) for x in circumstances_list]):
#                for i in range(j+1,len(lines)):
#                    print(list(lines[i]))
#                    # newf.write(lines[i])
#                j=i
#         else:
#             j=j+1
# #         # if j==len(lines):
# #         #     break
# #     # print(j)



#####rename files with the titles

# titles = [line.rstrip('\n') for line in open('title_gender_equality2(45).txt')]
# ids = [line.rstrip('\n') for line in open('itemid_gender_equality2(45).txt')]
# casedict=dictionary = dict(zip(ids, titles))
# print(casedict)
# # print(casedict['001-76766'])
# #
#
# directory="CaseLawGEcutfacts"
# for filename in os.listdir(directory):
#     # print(type(filename.rstrip('.txt')))
#     idfilename=filename.rstrip('.txt')
#     os.rename(directory+'\\'+filename, casedict[idfilename]+".txt")




###transformto txt
directorycut="CaseLawcomplete"
for filename in os.listdir(directorycut):
# document=Document('CaseLawPDF141\\001-60876.docx')
    try:
       document = Document(directorycut + "\\" + filename)
    except:
        continue
    file=filename.strip('docx')
    newf=open("CaseLawtxt\\"+file+"txt",'w',encoding='utf-8')
    for par in document.paragraphs:
                newf.write(par.text+"\n")
